from collections.abc import Generator
from pathlib import Path
from types import SimpleNamespace
from uuid import UUID

import fitz
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.auth import AuthenticatedUser, get_current_user
from app.db.base import Base
from app.db.session import get_db
from app.main import app
from app.models import Document, DocumentChunk, Reference, ThesisProject, UserProfile
from app.services.chunking import chunk_text


@pytest.fixture()
def db() -> Generator[Session, None, None]:
    engine = create_engine("sqlite+pysqlite:///:memory:", connect_args={"check_same_thread": False}, poolclass=StaticPool)
    Base.metadata.create_all(engine)
    TestingSessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

    with TestingSessionLocal() as session:
        yield session

    Base.metadata.drop_all(engine)


@pytest.fixture(autouse=True)
def clear_dependency_overrides() -> Generator[None, None, None]:
    app.dependency_overrides.clear()
    yield
    app.dependency_overrides.clear()


def client_for(db: Session, auth_user_id: str = "auth-owner") -> TestClient:
    def override_get_db() -> Generator[Session, None, None]:
        yield db

    app.dependency_overrides[get_db] = override_get_db
    app.dependency_overrides[get_current_user] = lambda: AuthenticatedUser(
        auth_user_id=auth_user_id,
        email=f"{auth_user_id}@example.com",
    )
    return TestClient(app)


def seed_documents(db: Session) -> dict[str, object]:
    owner = UserProfile(auth_user_id="auth-owner", email="owner@example.com")
    other = UserProfile(auth_user_id="auth-other", email="other@example.com")
    db.add_all([owner, other])
    db.flush()

    project = ThesisProject(owner_id=owner.id, title="Owned project")
    other_project = ThesisProject(owner_id=other.id, title="Other project")
    db.add_all([project, other_project])
    db.flush()

    document = Document(
        project_id=project.id,
        filename="draft.txt",
        content_type="text/plain",
        document_type="thesis_draft",
        size_bytes=128,
    )
    other_document = Document(project_id=other_project.id, filename="other.txt")
    db.add_all([document, other_document])
    db.flush()

    chunk = DocumentChunk(document_id=document.id, chunk_index=0, content="Chunk text")
    db.add(chunk)
    db.commit()
    db.refresh(project)
    db.refresh(document)
    db.refresh(other_document)
    db.refresh(chunk)

    return {
        "project": project,
        "document": document,
        "other_document": other_document,
        "chunk": chunk,
    }


def test_list_project_documents_for_owner(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)

    response = client.get(f"/api/v1/projects/{records['project'].id}/documents")

    assert response.status_code == 200
    assert response.json()[0]["id"] == str(records["document"].id)
    assert response.json()[0]["filename"] == "draft.txt"
    assert response.json()[0]["status"] == "uploaded"
    assert response.json()[0]["parse_status"] == "pending"


def test_create_text_document_saves_parsed_document_and_chunks(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)
    raw_text = "\n\n".join(
        [
            "Introduction\nThis thesis studies multi-agent review workflows.",
            "Methodology\nThe system chunks pasted thesis text before analysis.",
        ]
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents/text",
        json={
            "document_type": "thesis_draft",
            "title": "Methodology section",
            "raw_text": raw_text,
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["project_id"] == str(records["project"].id)
    assert body["document_type"] == "thesis_draft"
    assert body["title"] == "Methodology section"
    assert body["parse_status"] == "parsed"
    assert body["word_count"] == len(raw_text.split())
    assert body["chunk_count"] == 1

    document = db.scalar(select(Document).where(Document.id == UUID(body["id"])))
    chunks = list(
        db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == UUID(body["id"]))
            .order_by(DocumentChunk.chunk_index)
        )
    )
    assert document is not None
    assert document.filename == "Methodology section"
    assert document.raw_text == raw_text
    assert document.status == "parsed"
    assert document.parse_status == "parsed"
    assert document.content_type == "text/plain"
    assert document.size_bytes == len(raw_text.encode("utf-8"))
    assert [chunk.chunk_index for chunk in chunks] == [0]
    assert chunks[0].content.startswith("Introduction")


def test_create_text_document_splits_long_text_into_ordered_chunks(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)
    raw_text = " ".join(f"word{i}" for i in range(720))

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents/text",
        json={
            "document_type": "thesis_draft",
            "title": "Long draft",
            "raw_text": raw_text,
        },
    )

    assert response.status_code == 201
    assert response.json()["word_count"] == 720
    assert response.json()["chunk_count"] == 3
    chunks = list(
        db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == UUID(response.json()["id"]))
            .order_by(DocumentChunk.chunk_index)
        )
    )
    assert [chunk.chunk_index for chunk in chunks] == [0, 1, 2]
    assert len(chunks[0].content.split()) == 350
    assert len(chunks[1].content.split()) == 350
    assert len(chunks[2].content.split()) == 20


def test_chunk_text_preserves_section_heading_across_splits() -> None:
    raw_text = "\n".join(["Methodology", " ".join(f"token{i}" for i in range(12))])

    chunks = chunk_text(raw_text, max_words=8)

    assert [chunk.chunk_index for chunk in chunks] == [0, 1]
    assert chunks[0].content.startswith("Methodology")
    assert chunks[1].content.startswith("Methodology")
    assert len(chunks[0].content.split()) <= 8
    assert len(chunks[1].content.split()) <= 8


def test_chunk_text_handles_empty_text_safely() -> None:
    assert chunk_text("") == []
    assert chunk_text(" \n\n\t ") == []


def test_create_text_document_validates_project_ownership(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db, "auth-other")

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents/text",
        json={
            "document_type": "thesis_draft",
            "title": "Blocked",
            "raw_text": "This should not be stored.",
        },
    )

    assert response.status_code == 404
    assert db.scalar(select(Document).where(Document.filename == "Blocked")) is None


def test_upload_document_stores_metadata_and_file(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("draft.txt", b"Chapter 1\nResearch text", "text/plain")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["project_id"] == str(records["project"].id)
    assert body["filename"] == "draft.txt"
    assert body["document_type"] == "thesis_draft"
    assert body["content_type"] == "text/plain"
    assert body["size_bytes"] == len(b"Chapter 1\nResearch text")
    assert body["status"] == "uploaded"
    assert body["parse_status"] == "pending"
    assert body["raw_text"] is None
    assert body["parse_metadata"] is None

    storage_path = Path(body["storage_path"])
    assert storage_path.exists()
    assert storage_path.read_bytes() == b"Chapter 1\nResearch text"

    document = db.scalar(select(Document).where(Document.id == UUID(body["id"])))
    assert document is not None
    assert document.storage_path == body["storage_path"]


def test_upload_pdf_extracts_text_and_creates_chunks(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    pdf_bytes = build_pdf_bytes("Introduction\nThis thesis evaluates agent review workflows.")
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("draft.pdf", pdf_bytes, "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "draft.pdf"
    assert body["status"] == "parsed"
    assert body["parse_status"] == "parsed"
    assert "agent review workflows" in body["raw_text"]

    chunks = list(
        db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == UUID(body["id"]))
            .order_by(DocumentChunk.chunk_index)
        )
    )
    assert len(chunks) == 1
    assert "agent review workflows" in chunks[0].content


def test_upload_empty_pdf_marks_parse_failed(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    pdf_bytes = build_pdf_bytes("")
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("scanned.pdf", pdf_bytes, "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["raw_text"] is None
    assert db.scalar(select(DocumentChunk).where(DocumentChunk.document_id == UUID(body["id"]))) is None


def test_upload_corrupt_pdf_marks_parse_failed(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("broken.pdf", b"not a pdf", "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["raw_text"] is None


def test_upload_docx_extracts_paragraph_text_and_creates_chunks(
    db: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    records = seed_documents(db)
    client = client_for(db)
    docx_bytes = build_docx_bytes(["Introduction", "This thesis evaluates agent review workflows."])
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={
            "file": (
                "draft.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "draft.docx"
    assert body["status"] == "parsed"
    assert body["parse_status"] == "parsed"
    assert "Introduction" in body["raw_text"]
    assert "agent review workflows" in body["raw_text"]

    chunks = list(
        db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == UUID(body["id"]))
            .order_by(DocumentChunk.chunk_index)
        )
    )
    assert len(chunks) == 1
    assert "agent review workflows" in chunks[0].content


def test_upload_empty_docx_marks_parse_failed(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    docx_bytes = build_docx_bytes([])
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={
            "file": (
                "empty.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["raw_text"] is None
    assert db.scalar(select(DocumentChunk).where(DocumentChunk.document_id == UUID(body["id"]))) is None


def test_upload_malformed_docx_marks_parse_failed(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={
            "file": (
                "broken.docx",
                b"not a docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["raw_text"] is None


def test_upload_bibtex_extracts_references(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    bibtex = b"""@article{smith2024agents,
  title = {Agent Workflows for Thesis Review},
  author = {Smith, Ada and Khan, Omar},
  year = {2024},
  journal = {Journal of Research Automation},
  doi = {10.1234/example},
  url = {https://example.com/paper}
}
"""
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "reference_file"},
        files={"file": ("refs.bib", bibtex, "text/plain")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "refs.bib"
    assert body["status"] == "parsed"
    assert body["parse_status"] == "parsed"
    assert "smith2024agents" in body["raw_text"]

    reference = db.scalar(select(Reference).where(Reference.document_id == UUID(body["id"])))
    assert reference is not None
    assert reference.project_id == records["project"].id
    assert reference.citation_key == "smith2024agents"
    assert reference.title == "Agent Workflows for Thesis Review"
    assert reference.authors == ["Smith, Ada", "Khan, Omar"]
    assert reference.year == 2024
    assert reference.venue == "Journal of Research Automation"
    assert reference.doi == "10.1234/example"
    assert reference.url == "https://example.com/paper"
    assert reference.raw_bibtex.startswith("@article{smith2024agents")


def test_upload_malformed_bibtex_marks_parse_failed(
    db: Session,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "reference_file"},
        files={"file": ("broken.bib", b"this is not bibtex", "text/plain")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["raw_text"] is None
    assert db.scalar(select(Reference).where(Reference.document_id == UUID(body["id"]))) is None


def test_upload_csv_saves_summary_metadata(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "result_file"},
        files={"file": ("results.csv", b"model,accuracy\nbaseline,0.82\nagent,0.91\n", "text/csv")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "results.csv"
    assert body["status"] == "parsed"
    assert body["parse_status"] == "parsed"
    assert body["parse_metadata"] == {
        "row_count": 2,
        "column_names": ["model", "accuracy"],
        "sample_rows": [
            {"model": "baseline", "accuracy": 0.82},
            {"model": "agent", "accuracy": 0.91},
        ],
    }

    document = db.scalar(select(Document).where(Document.id == UUID(body["id"])))
    assert document is not None
    assert document.parse_metadata == body["parse_metadata"]


def test_upload_empty_csv_marks_parse_failed(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db)
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "result_file"},
        files={"file": ("empty.csv", b"", "text/csv")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "parse_failed"
    assert body["parse_status"] == "failed"
    assert body["parse_metadata"] is None


def test_upload_document_rejects_unsupported_type(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("draft.exe", b"not allowed", "application/octet-stream")},
    )

    assert response.status_code == 400
    assert db.scalar(select(Document).where(Document.filename == "draft.exe")) is None


def test_upload_document_rejects_large_file(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("large.txt", b"x" * (25 * 1024 * 1024 + 1), "text/plain")},
    )

    assert response.status_code == 413
    assert db.scalar(select(Document).where(Document.filename == "large.txt")) is None


def test_upload_document_validates_project_ownership(db: Session, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    records = seed_documents(db)
    client = client_for(db, "auth-other")
    monkeypatch.setattr(
        "app.services.documents.get_settings",
        lambda: SimpleNamespace(upload_storage_dir=str(tmp_path)),
    )

    response = client.post(
        f"/api/v1/projects/{records['project'].id}/documents",
        data={"document_type": "thesis_draft"},
        files={"file": ("blocked.txt", b"This should not be stored.", "text/plain")},
    )

    assert response.status_code == 404
    assert db.scalar(select(Document).where(Document.filename == "blocked.txt")) is None
    assert not any(tmp_path.iterdir())


def test_get_document_hides_other_users_document(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)

    owned_response = client.get(f"/api/v1/documents/{records['document'].id}")
    other_response = client.get(f"/api/v1/documents/{records['other_document'].id}")

    assert owned_response.status_code == 200
    assert owned_response.json()["id"] == str(records["document"].id)
    assert other_response.status_code == 404


def test_delete_document_removes_related_chunks(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db)

    response = client.delete(f"/api/v1/documents/{records['document'].id}")
    document = db.scalar(select(Document).where(Document.id == records["document"].id))
    chunk = db.scalar(select(DocumentChunk).where(DocumentChunk.id == records["chunk"].id))

    assert response.status_code == 204
    assert document is None
    assert chunk is None


def test_list_documents_hides_other_users_project(db: Session) -> None:
    records = seed_documents(db)
    client = client_for(db, "auth-other")

    response = client.get(f"/api/v1/projects/{records['project'].id}/documents")

    assert response.status_code == 404


def build_pdf_bytes(text: str) -> bytes:
    pdf = fitz.open()
    page = pdf.new_page()
    if text:
        page.insert_text((72, 72), text)
    return pdf.tobytes()


def build_docx_bytes(paragraphs: list[str]) -> bytes:
    from io import BytesIO

    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
